# 2023 which has sol. in question paper
# 
# import re
# import json
# from docx import Document

# def extract_questions_from_docx(docx_path):
#     # Load document
#     doc = Document(docx_path)
#     text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

#     # Remove page markers like "- 2 -", "NEET (UG)-2022 (Code-Q1)" etc.
#     text = re.sub(r"-\s*\d+\s*-", "", text)
#     text = re.sub(r"NEET.*?\n", "", text)

#     # Normalize spacing
#     text = re.sub(r"\n+", "\n", text).strip()

#     # Find all question number positions
#     qnum_pattern = re.compile(r"(?<!\d)(\d+)\.\s")
#     matches = list(qnum_pattern.finditer(text))

#     questions = []

#     for i, match in enumerate(matches):
#         qnum = int(match.group(1))
#         start_idx = match.end()

#         if i + 1 < len(matches):
#             next_match = matches[i + 1]
#             end_idx = next_match.start()
#         else:
#             end_idx = len(text)

#         block = text[match.start():end_idx].strip()

#         # Extract everything before 'Sol.'
#         sol_split = re.split(r"\bSol\.\b", block, maxsplit=1, flags=re.IGNORECASE)
#         question_text_block = sol_split[0].strip()

#         # Extract options (1), (2), (3), (4)
#         options = re.findall(r"\(\d+\)\s*([^\(]+?)(?=\s*\(\d+\)|$)", question_text_block, re.DOTALL)
#         options = [opt.strip() for opt in options]

#         # Extract question text (before first option)
#         qtext_split = re.split(r"\(\d+\)", question_text_block, maxsplit=1)
#         question_text = qtext_split[0].strip() if qtext_split else question_text_block

#         # Extract answer if present inside question text (optional)
#         answer_match = re.search(r"Answer\s*\(?(\d+)\)?", question_text_block, re.IGNORECASE)
#         answer = answer_match.group(1) if answer_match else None

#         # Build question dictionary
#         question_data = {
#             "question_number": qnum,  # will convert to string later
#             "question_text": question_text,
#             "question_diagram_description": "null",
#             "options": {
#                 "option_1": options[0] if len(options) > 0 else "",
#                 "option_2": options[1] if len(options) > 1 else "",
#                 "option_3": options[2] if len(options) > 2 else "",
#                 "option_4": options[3] if len(options) > 3 else "",
#             },
#             "options_diagram_description": "null",
#             "answer": answer
#         }

#         questions.append(question_data)

#     # --- Filter to keep only strictly consecutive question numbers and convert to string ---
#     filtered_questions = []
#     expected_qnum = 1
#     for q in questions:
#         if int(q["question_number"]) == expected_qnum:
#             q["question_number"] = str(expected_qnum)  # convert to string
#             filtered_questions.append(q)
#             expected_qnum += 1
#         else:
#             print(f"⚠️ Removing question {q['question_number']} (breaks sequence)")

#     return filtered_questions


# if __name__ == "__main__":
#     docx_path = r"D:\NEET\pdf\2023.docx"
#     data = extract_questions_from_docx(docx_path)

#     output_path = "no_diagram_questions_2023.json"
#     with open(output_path, "w", encoding="utf-8") as f:
#         json.dump(data, f, indent=4, ensure_ascii=False)

#     print(f"✅ Extracted {len(data)} questions → {output_path}")


# 2022_question paper - no solutions in wuestion paper

import re
import json
from docx import Document

def extract_questions_from_docx(docx_path):
    # Load document
    doc = Document(docx_path)
    text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

    # Remove page markers like "- 2 -", "NEET (UG)-2022 (Code-Q1)" etc.
    text = re.sub(r"-\s*\d+\s*-", "", text)
    text = re.sub(r"NEET.*?\n", "", text)

    # Split into questions — each starts with a number followed by a period (e.g., "1.")
    question_blocks = re.split(r"\n(?=\d+\.\s)", text)

    questions = []

    for block in question_blocks:
        block = block.strip()
        if not re.match(r"^\d+\.", block):
            continue

        # Extract question number
        qnum_match = re.match(r"^(\d+)\.\s*(.*)", block, re.DOTALL)
        if not qnum_match:
            continue

        qnum = qnum_match.group(1)
        rest = qnum_match.group(2)

        # Extract answer
        answer_match = re.search(r"Answer\s*\(?(\d+)\)?", rest, re.IGNORECASE)
        answer = answer_match.group(1) if answer_match else None

        # Extract options (1)...(4)
        options = re.findall(r"\(\d+\)\s*([^\(]+?)(?=\s*\(\d+\)|\s*Answer|$)", rest, re.DOTALL)
        options = [opt.strip() for opt in options]

        # Extract question text (everything before first option)
        qtext_match = re.split(r"\(\d+\)", rest, maxsplit=1)
        question_text = qtext_match[0].strip() if qtext_match else rest.strip()

        # Build question dictionary
        question_data = {
            "question_number": qnum,
            "question_text": question_text,
            "question_diagram_description": "null",
            "options": {
                "option_1": options[0] if len(options) > 0 else "",
                "option_2": options[1] if len(options) > 1 else "",
                "option_3": options[2] if len(options) > 2 else "",
                "option_4": options[3] if len(options) > 3 else "",
            },
            "options_diagram_description": "null",
            "answer": answer
        }

        questions.append(question_data)

    return questions


if __name__ == "__main__":
    docx_path = r"D:\NEET\pdf\2024.docx"
    data = extract_questions_from_docx(docx_path)

    with open("no_diagram_questions_2024.json", "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4, ensure_ascii=False)

    print(f"✅ Extracted {len(data)} questions to neet_questions.json")
