# import pytesseract
# from PIL import Image

# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# # Load image from the specified path
# image = Image.open(r'D:\\NEET\\extracted_images\\page_1.jpg')

# # Extract text using pytesseract
# text = pytesseract.image_to_string(image)

# print(text)

from docx import Document
import json

def extract_syllabus_to_json(docx_path):
    doc = Document(docx_path)
    subjects_data = []
    subject_name = None
    table_index = 0
    
    # Collect all tables first as a list
    all_tables = list(doc.tables)
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Detect subject header
        if text.startswith('NEET') and 'Syllabus' in text:
            # Save the previous subject info if exists
            if subject_name is not None and table_index < len(all_tables):
                units = []
                table = all_tables[table_index]
                for row in table.rows[1:]:  # skip header
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 3:
                        units.append({
                            'Unit': cells[0],
                            'Topic': cells[1],
                            'Subtopic': cells[2]
                        })
                subjects_data.append({
                    'Subject': subject_name,
                    'Units': units
                })
                table_index += 1
            
            # Set new subject
            subject_name = text.replace('NEET ','').replace(' Syllabus','')
    
    # Handle the last subject and table
    if subject_name is not None and table_index < len(all_tables):
        units = []
        table = all_tables[table_index]
        for row in table.rows[1:]:  # skip header
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 3:
                units.append({
                    'Unit': cells[0],
                    'Topic': cells[1],
                    'Subtopic': cells[2]
                })
        subjects_data.append({
            'Subject': subject_name,
            'Units': units
        })

    return subjects_data

# Usage
docx_path = r'D:\NEET\syllabus_spellchecked.docx'
syllabus_json = extract_syllabus_to_json(docx_path)

# Save the results to JSON file
with open('neet_syllabus_full_structured.json', 'w', encoding='utf-8') as f:
    json.dump(syllabus_json, f, indent=2)

print("JSON extraction complete. Check 'neet_syllabus_full_structured.json'")
